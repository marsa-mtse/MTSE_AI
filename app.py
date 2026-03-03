# ==========================================================
# MTSE AI — FULL ENTERPRISE MASTER BUILD (FINAL STABLE)
# ==========================================================

import streamlit as st
from groq import Groq
import pandas as pd
import numpy as np
import sqlite3
import hashlib
import datetime
import io
import zipfile
import pdfplumber
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

# ==========================================================
# PAGE CONFIG
# ==========================================================

st.set_page_config(page_title="MTSE AI", page_icon="🚀", layout="wide")

# ==========================================================
# STYLE
# ==========================================================

st.markdown("""
<style>
body {
    background: linear-gradient(135deg,#0f172a,#111827);
    color:white;
}
section[data-testid="stSidebar"] {
    background:#0b1120;
}
.stButton > button {
    width:100%;
    border-radius:8px;
    height:42px;
    font-weight:600;
    background:rgba(255,255,255,0.05);
    color:white;
}
.stButton > button:hover {
    background:linear-gradient(90deg,#00c6ff,#0072ff);
}
</style>
""", unsafe_allow_html=True)

# ==========================================================
# LANGUAGE
# ==========================================================

if "lang" not in st.session_state:
    st.session_state.lang = "ar"

def t(ar, en):
    return ar if st.session_state.lang == "ar" else en

# ==========================================================
# DATABASE
# ==========================================================

conn = sqlite3.connect("mtse.db", check_same_thread=False)
c = conn.cursor()

c.execute("""
CREATE TABLE IF NOT EXISTS users(
id INTEGER PRIMARY KEY AUTOINCREMENT,
username TEXT UNIQUE,
password TEXT,
role TEXT,
package TEXT,
credits INTEGER
)
""")

c.execute("""
CREATE TABLE IF NOT EXISTS cost_history(
id INTEGER PRIMARY KEY AUTOINCREMENT,
project TEXT,
total REAL,
created_at TEXT
)
""")

conn.commit()

def hash_pass(p):
    return hashlib.sha256(p.encode()).hexdigest()

# Admin default
c.execute("SELECT * FROM users WHERE username='admin'")
if not c.fetchone():
    c.execute("INSERT INTO users VALUES(NULL,?,?,?,?,?)",
              ("admin",hash_pass("admin123"),"admin","enterprise",9999))
    conn.commit()

# ==========================================================
# AUTH
# ==========================================================

if "user" not in st.session_state:
    st.session_state.user = None

def login(u,p):
    c.execute("SELECT * FROM users WHERE username=? AND password=?",
              (u,hash_pass(p)))
    return c.fetchone()

# ==========================================================
# AI SETUP
# ==========================================================

client = Groq(api_key=st.secrets["GROQ_API_KEY"])

def ask_ai(system,user):
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role":"system","content":system},
            {"role":"user","content":user}
        ],
        temperature=0.5,
        max_tokens=2000
    )
    return response.choices[0].message.content

def classify(text):
    return ask_ai("أنت مصنف محتوى محترف",
                  f"حدد نوع المحتوى: هندسي / مالي / إداري / تسويقي / قانوني / مختلط\n{text}")

# ==========================================================
# COST ENGINE
# ==========================================================

def cost_engine(base):
    return {
        "Base Cost": base,
        "Indirect 10%": base*0.10,
        "Waste 5%": base*0.05,
        "Admin 7%": base*0.07,
        "Profit 15%": base*0.15,
        "Conservative": base*1.05,
        "Moderate": base*1.15,
        "Aggressive": base*1.25
    }

# ==========================================================
# PDF REPORT
# ==========================================================

def generate_pdf(content):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    elements = []
    styles = getSampleStyleSheet()
    elements.append(Paragraph("MTSE AI Professional Report",styles["Heading1"]))
    elements.append(Spacer(1,0.5*inch))
    elements.append(Paragraph(content,styles["Normal"]))
    doc.build(elements)
    buffer.seek(0)
    return buffer

# ==========================================================
# LOGIN PAGE
# ==========================================================

if not st.session_state.user:

    st.title(t("تسجيل الدخول","Login"))

    u = st.text_input(t("اسم المستخدم","Username"))
    p = st.text_input(t("كلمة المرور","Password"),type="password")

    if st.button(t("دخول","Login")):
        user = login(u,p)
        if user:
            st.session_state.user = user
            st.rerun()
        else:
            st.error(t("بيانات غير صحيحة","Invalid credentials"))

# ==========================================================
# MAIN SYSTEM
# ==========================================================

else:

    user = st.session_state.user
    role = user[3]
    credits = user[5]

    with st.sidebar:

        st.markdown("## 🚀 MTSE AI")

        if st.button("🌍 العربية / English"):
            st.session_state.lang = "en" if st.session_state.lang=="ar" else "ar"
            st.rerun()

        page = st.radio(
            t("القائمة","Navigation"),
            [
                t("المحلل الشامل","Universal Analyzer"),
                t("محرك المقايسات","Cost Engine"),
                t("تحليل السوشيال","Social Media"),
                t("التقارير","Reports"),
                t("الإدارة","Admin"),
                t("تسجيل الخروج","Logout")
            ]
        )

        st.write(f"👤 {user[1]}")
        st.write(f"⚡ {t('الرصيد','Credits')}: {credits}")

    # ==========================================================
    # UNIVERSAL ANALYZER
    # ==========================================================

    if page == t("المحلل الشامل","Universal Analyzer"):

        text = st.text_area(t("اكتب نص","Enter text"))
        file = st.file_uploader(t("ارفع ملف","Upload file"))

        if st.button(t("تحليل","Analyze")):

            if credits <= 0:
                st.error(t("لا يوجد رصيد","No credits left"))
            else:

                combined = text if text else ""

                if file:
                    if file.name.endswith(".pdf"):
                        with pdfplumber.open(file) as pdf:
                            for pg in pdf.pages:
                                combined += pg.extract_text() or ""
                    elif file.name.endswith(".docx"):
                        doc = Document(file)
                        for para in doc.paragraphs:
                            combined += para.text
                    elif file.name.endswith(".csv"):
                        df = pd.read_csv(file)
                        combined += df.to_string()
                    elif file.name.endswith(".xlsx"):
                        df = pd.read_excel(file)
                        combined += df.to_string()
                    elif file.name.endswith(".zip"):
                        with zipfile.ZipFile(file) as z:
                            combined += "ZIP contains: " + ", ".join(z.namelist())
                    else:
                        combined += file.read().decode("utf-8",errors="ignore")

                content_type = classify(combined)
                st.subheader(t("نوع المحتوى","Content Type"))
                st.write(content_type)

                result = ask_ai(
                    "أنت محلل مشاريع شامل يقدم تقرير احترافي شامل",
                    combined
                )

                st.markdown(result)

                pdf = generate_pdf(result)
                st.download_button(t("تحميل PDF","Download PDF"),pdf,"MTSE_Report.pdf")

                c.execute("UPDATE users SET credits=credits-1 WHERE username=?",(user[1],))
                conn.commit()

    # ==========================================================
    # COST ENGINE
    # ==========================================================

    elif page == t("محرك المقايسات","Cost Engine"):

        qty = st.number_input(t("الكمية","Quantity"),0.0)
        price = st.number_input(t("سعر الوحدة","Unit Price"),0.0)

        if st.button(t("احسب","Calculate")):
            base = qty*price
            result = cost_engine(base)
            st.json(result)

            c.execute("INSERT INTO cost_history VALUES(NULL,?,?,?)",
                      ("Manual",base,str(datetime.datetime.now())))
            conn.commit()

    # ==========================================================
    # SOCIAL MEDIA
    # ==========================================================

    elif page == t("تحليل السوشيال","Social Media"):

        content = st.text_area(t("ضع الرابط أو المحتوى","Paste link or content"))

        if st.button(t("تحليل سوشيال","Analyze Social")):
            result = ask_ai(
                "أنت استراتيجي تسويق محترف يقدم تحليل أداء + خطة تطوير",
                content
            )
            st.markdown(result)

    # ==========================================================
    # REPORTS
    # ==========================================================

    elif page == t("التقارير","Reports"):

        final = st.text_area(t("محتوى التقرير","Report content"))

        if st.button(t("إنشاء تقرير","Generate Report")):
            pdf = generate_pdf(final)
            st.download_button(t("تحميل التقرير","Download Report"),
                               pdf,"MTSE_Final_Report.pdf")

    # ==========================================================
    # ADMIN
    # ==========================================================

    elif page == t("الإدارة","Admin"):

        if role != "admin":
            st.error(t("صلاحية غير متاحة","Admin only"))
        else:

            new_user = st.text_input(t("اسم مستخدم جديد","New username"))
            new_pass = st.text_input(t("كلمة المرور","Password"),type="password")
            package = st.selectbox(t("الباقة","Package"),
                                   ["free","pro","enterprise"])
            credits_new = st.number_input(t("الرصيد","Credits"),0)

            if st.button(t("إنشاء مستخدم","Create User")):
                c.execute("INSERT INTO users VALUES(NULL,?,?,?,?,?)",
                          (new_user,hash_pass(new_pass),
                           "client",package,credits_new))
                conn.commit()
                st.success(t("تم الإنشاء","Created"))

            df = pd.read_sql("SELECT username,role,package,credits FROM users",conn)
            st.dataframe(df)

    # ==========================================================
    # LOGOUT
    # ==========================================================

    elif page == t("تسجيل الخروج","Logout"):
        st.session_state.user=None
        st.rerun()
# ==============================
# IMPORTS (AI + FILES)
# ==============================

from groq import Groq
import pandas as pd
import numpy as np
import pdfplumber
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
import io
import zipfile

# ==============================
# GROQ INIT
# ==============================

groq_key = st.secrets.get("GROQ_API_KEY", None)
client = Groq(api_key=groq_key) if groq_key else None

def ask_ai(system_prompt, user_prompt):
    if not client:
        return "Groq not configured."
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role":"system","content":system_prompt},
            {"role":"user","content":user_prompt}
        ],
        temperature=0.4,
        max_tokens=2000
    )
    return response.choices[0].message.content

def classify_content(text):
    return ask_ai(
        "أنت مصنف محتوى محترف",
        f"حدد نوع المحتوى: هندسي / مالي / إداري / تسويقي / قانوني / مختلط\n{text[:4000]}"
    )

# ==============================
# ENTERPRISE PDF
# ==============================

def generate_enterprise_pdf(content, username):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("MTSE AI Enterprise Report", styles["Heading1"]))
    elements.append(Spacer(1,0.5*inch))
    elements.append(Paragraph(f"Client: {username}", styles["Normal"]))
    elements.append(Spacer(1,0.5*inch))
    elements.append(Paragraph(content.replace("\n","<br/>"), styles["Normal"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer

# ==============================
# ADVANCED COST ENGINE
# ==============================

def advanced_cost_engine(qty, price):
    base = qty * price
    indirect = base * 0.1
    waste = base * 0.05
    admin = base * 0.07
    profit = base * 0.15

    conservative = base * 1.05
    moderate = base * 1.15
    aggressive = base * 1.25

    return {
        "Base": base,
        "Indirect": indirect,
        "Waste": waste,
        "Admin": admin,
        "Profit": profit,
        "Conservative": conservative,
        "Moderate": moderate,
        "Aggressive": aggressive
    }

# ==============================
# EXTEND NAVIGATION
# ==============================

if st.session_state.user:

    with st.sidebar:
        extended_page = st.radio(
            t("محركات النظام","System Engines"),
            [
                t("المحلل الشامل","Universal Analyzer"),
                t("محرك المقايسات","Cost Engine"),
                t("تحليل السوشيال","Social Engine"),
                t("التقارير الاحترافية","Enterprise Reports")
            ]
        )

    username = st.session_state.user[1]

    # ==========================================================
    # UNIVERSAL ANALYZER
    # ==========================================================

    if extended_page == t("المحلل الشامل","Universal Analyzer"):

        st.title(t("المحلل الشامل","Universal Analyzer"))

        text_input = st.text_area(t("اكتب نص","Enter text"))
        uploaded_file = st.file_uploader(t("ارفع ملف","Upload file"))

        if st.button(t("تحليل","Analyze")):

            c.execute("SELECT credits_balance FROM users WHERE username=?", (username,))
            credits = c.fetchone()[0]

            if credits <= 0:
                st.error(t("لا يوجد رصيد","No credits left"))
            else:
                combined = text_input if text_input else ""

                if uploaded_file:
                    if uploaded_file.name.endswith(".pdf"):
                        with pdfplumber.open(uploaded_file) as pdf:
                            for page in pdf.pages:
                                combined += page.extract_text() or ""
                    elif uploaded_file.name.endswith(".docx"):
                        doc = Document(uploaded_file)
                        for para in doc.paragraphs:
                            combined += para.text
                    elif uploaded_file.name.endswith(".csv"):
                        df = pd.read_csv(uploaded_file)
                        combined += df.to_string()
                    elif uploaded_file.name.endswith(".xlsx"):
                        df = pd.read_excel(uploaded_file)
                        combined += df.to_string()
                    elif uploaded_file.name.endswith(".zip"):
                        with zipfile.ZipFile(uploaded_file) as z:
                            combined += "ZIP contains: " + ", ".join(z.namelist())
                    else:
                        combined += uploaded_file.read().decode("utf-8", errors="ignore")

                content_type = classify_content(combined)
                st.subheader(t("نوع المحتوى","Content Type"))
                st.write(content_type)

                result = ask_ai(
                    "أنت محلل مشاريع شامل يقدم تقرير احترافي يشمل ملخص تنفيذي وتحليل وتوصيات",
                    combined[:12000]
                )

                st.markdown(result)

                # Save analysis
                c.execute("""
                INSERT INTO analyses VALUES(NULL,?,?,?,?,?)
                """, (
                    username,
                    "universal",
                    combined[:500],
                    result,
                    datetime.datetime.now().isoformat()
                ))
                c.execute("""
                UPDATE users SET credits_balance = credits_balance - 1
                WHERE username=?
                """, (username,))
                conn.commit()

                pdf = generate_enterprise_pdf(result, username)
                st.download_button(t("تحميل PDF","Download PDF"),
                                   pdf,
                                   "MTSE_Enterprise_Report.pdf")

    # ==========================================================
    # COST ENGINE
    # ==========================================================

    elif extended_page == t("محرك المقايسات","Cost Engine"):

        st.title(t("محرك المقايسات","Cost Engine"))

        qty = st.number_input(t("الكمية","Quantity"), 0.0)
        price = st.number_input(t("سعر الوحدة","Unit Price"), 0.0)

        if st.button(t("احسب","Calculate")):
            result = advanced_cost_engine(qty, price)
            st.json(result)

            c.execute("""
            INSERT INTO analyses VALUES(NULL,?,?,?,?,?)
            """, (
                username,
                "cost",
                f"Qty:{qty} Price:{price}",
                str(result),
                datetime.datetime.now().isoformat()
            ))
            conn.commit()

    # ==========================================================
    # SOCIAL ENGINE
    # ==========================================================

    elif extended_page == t("تحليل السوشيال","Social Engine"):

        st.title(t("تحليل السوشيال","Social Engine"))

        social_input = st.text_area(t("ضع الرابط أو المحتوى","Paste link or content"))

        if st.button(t("تحليل","Analyze")):

            result = ask_ai(
                "أنت استراتيجي تسويق محترف. قدم تحليل أداء + تحليل هوية + خطة 30 يوم + توصيات تطوير.",
                social_input
            )

            st.markdown(result)

            c.execute("""
            INSERT INTO analyses VALUES(NULL,?,?,?,?,?)
            """, (
                username,
                "social",
                social_input[:500],
                result,
                datetime.datetime.now().isoformat()
            ))
            conn.commit()

    # ==========================================================
    # ENTERPRISE REPORT GENERATOR
    # ==========================================================

    elif extended_page == t("التقارير الاحترافية","Enterprise Reports"):

        st.title(t("التقارير الاحترافية","Enterprise Reports"))

        report_text = st.text_area(t("محتوى التقرير","Report Content"))

        if st.button(t("إنشاء تقرير","Generate Report")):
            pdf = generate_enterprise_pdf(report_text, username)

            c.execute("""
            INSERT INTO analyses VALUES(NULL,?,?,?,?,?)
            """, (
                username,
                "report",
                report_text[:500],
                report_text,
                datetime.datetime.now().isoformat()
            ))
            conn.commit()

            st.download_button(t("تحميل التقرير","Download Report"),
                               pdf,
                               "MTSE_Final_Report.pdf")
            # ==========================================================
# PART 3 — ENTERPRISE HARDENING LAYER
# ==========================================================

import re

# ==========================================================
# ACTIVITY LOG TABLE
# ==========================================================

c.execute("""
CREATE TABLE IF NOT EXISTS activity_log(
id INTEGER PRIMARY KEY AUTOINCREMENT,
username TEXT,
action TEXT,
details TEXT,
created_at TEXT
)
""")
conn.commit()

def log_activity(username, action, details=""):
    c.execute("""
    INSERT INTO activity_log VALUES(NULL,?,?,?,?)
    """, (
        username,
        action,
        details,
        datetime.datetime.now().isoformat()
    ))
    conn.commit()

# ==========================================================
# PASSWORD POLICY
# ==========================================================

def validate_password(password):
    if len(password) < 8:
        return False
    return True

# ==========================================================
# LOGIN RATE LIMITER
# ==========================================================

if "login_attempts" not in st.session_state:
    st.session_state.login_attempts = 0

def login_guard():
    if st.session_state.login_attempts >= 5:
        st.error("Too many login attempts. Try again later.")
        st.stop()

# Apply login guard
if not st.session_state.user:
    login_guard()

# ==========================================================
# SMART CREDIT ENGINE
# ==========================================================

CREDIT_COST = {
    "universal": 2,
    "social": 2,
    "cost": 1,
    "report": 3
}

def charge_credits(username, action_type, extra=0):
    cost = CREDIT_COST.get(action_type, 1) + extra
    c.execute("SELECT credits_balance FROM users WHERE username=?", (username,))
    balance = c.fetchone()[0]

    if balance < cost:
        return False

    c.execute("""
    UPDATE users SET credits_balance = credits_balance - ?
    WHERE username=?
    """, (cost, username))
    conn.commit()
    return True

# ==========================================================
# FILE SECURITY LAYER
# ==========================================================

ALLOWED_EXTENSIONS = ["pdf","docx","csv","xlsx","txt","zip"]
MAX_FILE_SIZE_MB = 10

def validate_file(uploaded_file):
    if uploaded_file is None:
        return True

    ext = uploaded_file.name.split(".")[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        st.error("File type not allowed.")
        return False

    size_mb = uploaded_file.size / (1024*1024)
    if size_mb > MAX_FILE_SIZE_MB:
        st.error("File too large.")
        return False

    return True

# ==========================================================
# INPUT SANITIZER
# ==========================================================

def sanitize_input(text):
    if not text:
        return ""
    text = re.sub(r"[<>]", "", text)
    return text[:15000]

# ==========================================================
# STRIPE SESSION VERIFICATION
# ==========================================================

def verify_payment(session_id, username):
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid":
            amount = session.amount_total / 100

            # Upgrade plan based on amount
            if amount == PLAN_PRICING["pro"]:
                plan = "pro"
            elif amount == PLAN_PRICING["enterprise"]:
                plan = "enterprise"
            else:
                return False

            credits = PLAN_CREDITS[plan]

            c.execute("""
            UPDATE users 
            SET package=?, credits_limit=?, credits_balance=?, billing_status='active'
            WHERE username=?
            """, (plan, credits, credits, username))

            c.execute("""
            UPDATE payments SET status='paid'
            WHERE stripe_session=?
            """, (session_id,))

            conn.commit()
            log_activity(username, "payment_verified", plan)
            return True

    except Exception:
        return False

# ==========================================================
# APPLY PAYMENT VERIFICATION IF SUCCESS PARAM EXISTS
# ==========================================================

query_params = st.query_params

if "payment" in query_params and query_params["payment"] == "success":
    c.execute("""
    SELECT stripe_session FROM payments 
    WHERE username=? AND status='pending'
    ORDER BY id DESC LIMIT 1
    """, (st.session_state.user[1],))
    last_payment = c.fetchone()

    if last_payment:
        verified = verify_payment(last_payment[0], st.session_state.user[1])
        if verified:
            st.success("Payment verified and plan upgraded.")
        else:
            st.error("Payment verification failed.")

# ==========================================================
# ROLE GUARD
# ==========================================================

def require_admin(role):
    if role != "admin":
        st.error("Access denied.")
        st.stop()

# ==========================================================
# ENHANCED ADMIN PANEL
# ==========================================================

if st.session_state.user:

    username = st.session_state.user[1]
    role = st.session_state.user[3]

    if role == "admin":

        st.markdown("---")
        st.header("Advanced Admin Controls")

        # Suspend user
        suspend_user = st.text_input("Suspend Username")
        if st.button("Suspend User"):
            c.execute("""
            UPDATE users SET billing_status='suspended'
            WHERE username=?
            """, (suspend_user,))
            conn.commit()
            log_activity(username, "suspend_user", suspend_user)
            st.success("User suspended.")

        # Adjust credits
        adjust_user = st.text_input("Adjust Credits Username")
        new_credits = st.number_input("New Credit Balance", 0)

        if st.button("Update Credits"):
            c.execute("""
            UPDATE users SET credits_balance=?
            WHERE username=?
            """, (new_credits, adjust_user))
            conn.commit()
            log_activity(username, "adjust_credits", adjust_user)
            st.success("Credits updated.")

        # View activity log
        if st.button("View Activity Logs"):
            logs = c.execute("""
            SELECT username, action, created_at 
            FROM activity_log ORDER BY id DESC LIMIT 50
            """).fetchall()
            for log in logs:
                st.write(log)
            
